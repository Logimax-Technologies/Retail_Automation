"""
generate_word_doc.py
Reads source Python test files and generates a professional Word document
(.docx) for manager review — Module-wise test coverage report.

Requires: pip install python-docx
"""

import os
import ast
import datetime
from pathlib import Path

# ── Module definitions: display name, folder, files, description ──────────────
MODULES = [
    {
        "folder": "Test_login",
        "display": "Login & Authentication",
        "files": ["Login.py"],
        "description": "Validates user login with valid credentials and verifies dashboard access."
    },
    {
        "folder": "Test_master",
        "display": "Master Data Management",
        "files": ["Metal.py", "Category.py", "Product.py", "Design.py",
                  "Subdesign.py", "Designmapping.py", "Subdesignmapping.py",
                  "MCVA.py", "StoneRateSettings.py"],
        "description": "Covers all master setup operations — metals, categories, products, designs, "
                       "sub-designs, mappings, making charge / value addition, and stone rate settings."
    },
    {
        "folder": "Test_lot",
        "display": "Lot Management",
        "files": ["Lot.py", "LotGenerate.py"],
        "description": "Tests lot creation, lot generation and assignment workflows."
    },
    {
        "folder": "Test_Tag",
        "display": "Tagging",
        "files": ["Tag.py"],
        "description": "Validates tag creation, tag scanning, lot-linking, and tag status updates."
    },
    {
        "folder": "Test_vendor",
        "display": "Vendor Management",
        "files": ["Vendor.py", "VendorApproval.py"],
        "description": "Covers vendor registration and approval workflows."
    },
    {
        "folder": "Test_Customer",
        "display": "Customer Order Management",
        "files": ["CustomerOrder.py", "KarigarAllotment.py"],
        "description": "Tests customer order creation, karigar allotment, and order tracking."
    },
    {
        "folder": "Test_EST",
        "display": "Estimation",
        "files": ["EST.py"],
        "description": "Validates estimation creation, tagging linkage, and estimation status updates."
    },
    {
        "folder": "Test_Bill",
        "display": "Billing & Receipts",
        "files": ["Bill.py", "BillingIssue.py", "BillingReceipt.py",
                  "BillingDenomination.py", "JewelNotDelivered.py",
                  "BillSplit.py", "SearchBill.py"],
        "description": "Full billing lifecycle — new bill creation (Sales, Purchase, Return, Order Delivery, "
                       "Repair Delivery), payment modes (Cash, Card, Cheque, Net Banking), receipt "
                       "management, denomination, jewel delivery status, and bill split."
    },
    {
        "folder": "Test_Purchase",
        "display": "Purchase Management",
        "files": ["PurchasePO.py", "GRNEntry.py", "SupplierBillEntry.py",
                  "HMIssueReceipt.py", "QCIssueReceipt.py", "PurchaseReturn.py",
                  "SmithSupplierPayment.py", "DebitCreditEntry.py",
                  "SmithMetalIssue.py", "RateFixGSTPurchase.py",
                  "SmithCompanyOpBal.py", "ApprovalRateFixing.py",
                  "ApprovalToInvoice.py"],
        "description": "End-to-end purchase process — PO creation, GRN entry, supplier bill, "
                       "HM/QC issue & receipt, purchase returns, smith payments, metal issue, "
                       "GST rate fixing, approval workflows, and opening balance."
    },
    {
        "folder": "Test_OldMetalProcess",
        "display": "Old Metal Process",
        "files": ["OldMetalProcess.py"],
        "description": "Validates old metal exchange entry and processing."
    },
    {
        "folder": "Test_StockIssue",
        "display": "Stock Issue",
        "files": ["StockIssue.py"],
        "description": "Tests stock issue to branches and internal departments."
    },
    {
        "folder": "Test_RepairOrder",
        "display": "Repair Order Management",
        "files": ["RepairOrder.py", "KarigarAllotment.py", "RepairOrderStatus.py"],
        "description": "Covers repair order creation, karigar allotment, and status updates."
    },
    {
        "folder": "Test_Inventory",
        "display": "Inventory Management",
        "files": ["BranchTransfer.py", "BranchTransferApproval.py",
                  "OrderLink.py", "TagUnlink.py", "NonTagReceipt.py"],
        "description": "Branch transfer requests and approvals, order linking/unlinking of tags, "
                       "and non-tag receipt processing."
    },
    {
        "folder": "Test_SectionTransfer",
        "display": "Section Transfer",
        "files": ["SectionTransfer.py"],
        "description": "Tests stock transfer between sections within the same branch."
    },
    {
        "folder": "Test_OtherInventory",
        "display": "Other Inventory (Packaging)",
        "files": ["InventoryCategory.py", "PackagingItemSize.py",
                  "OtherInventory.py", "ProductMapping.py",
                  "ProductPurchaseEntry.py", "PackagingItemIssue.py",
                  "OtherInventoryTagging.py"],
        "description": "Full packaging/other inventory lifecycle — categories, sizes, items, "
                       "product mapping, purchase entries, issue, and tagging."
    },
]


def _extract_functions_from_file(filepath: str):
    """Parse a Python file and return list of (function_name, docstring) tuples."""
    results = []
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            source = f.read()
        tree = ast.parse(source)
        for node in ast.walk(tree):
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                name = node.name
                doc = ast.get_docstring(node) or ""
                # Skip private helpers and dunder methods
                if not name.startswith("__"):
                    results.append((name, doc.strip()))
    except Exception as e:
        results.append(("(parse error)", str(e)))
    return results


def generate_word_doc():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return

    BASE_DIR = Path(__file__).parent
    OUTPUT_PATH = BASE_DIR / "Reports" / "Test_Coverage_Document.docx"
    OUTPUT_PATH.parent.mkdir(exist_ok=True)

    doc = Document()

    # ── Page Margins ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Helper to set cell background ───────────────────────────────────────────
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, color="CCCCCC"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # ── COVER PAGE ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AUTOMATION TEST COVERAGE DOCUMENT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)  # indigo

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Retail Automation Framework — Module-wise Test Case Report")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    now = datetime.datetime.now().strftime("%B %d, %Y  |  %H:%M")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {now}").font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # ── SUMMARY TABLE ────────────────────────────────────────────────────────────
    heading = doc.add_paragraph()
    heading.add_run("1.  Executive Summary").bold = True
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0x31, 0x2E, 0x81)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    headers = ["Module Group", "Files Covered", "Functions Found", "Status"]
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], "312E81")

    total_files = 0
    total_funcs = 0

    for mod in MODULES:
        row = tbl.add_row().cells
        folder_path = BASE_DIR / mod["folder"]
        file_count = 0
        func_count = 0
        for fname in mod["files"]:
            fpath = folder_path / fname
            if fpath.exists():
                file_count += 1
                funcs = _extract_functions_from_file(str(fpath))
                func_count += len([f for f in funcs if not f[0].startswith("_")])
        total_files += file_count
        total_funcs += func_count

        row[0].text = mod["display"]
        row[1].text = str(file_count)
        row[2].text = str(func_count)
        row[3].text = "✔ Automated"
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
        row[3].paragraphs[0].runs[0].bold = True
        for c in row:
            set_cell_border(c)

    # Totals row
    tot_row = tbl.add_row().cells
    tot_row[0].text = "TOTAL"
    tot_row[0].paragraphs[0].runs[0].bold = True
    tot_row[1].text = str(total_files)
    tot_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tot_row[1].paragraphs[0].runs[0].bold = True
    tot_row[2].text = str(total_funcs)
    tot_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tot_row[2].paragraphs[0].runs[0].bold = True
    tot_row[3].text = f"{len(MODULES)} Module Groups"
    tot_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for c in tot_row:
        set_cell_bg(c, "EEF2FF")
        set_cell_border(c)

    doc.add_paragraph()
    doc.add_page_break()

    # ── MODULE SECTIONS ──────────────────────────────────────────────────────────
    heading2 = doc.add_paragraph()
    heading2.add_run("2.  Detailed Module Coverage").bold = True
    heading2.runs[0].font.size = Pt(14)
    heading2.runs[0].font.color.rgb = RGBColor(0x31, 0x2E, 0x81)
    doc.add_paragraph()

    for idx, mod in enumerate(MODULES, 1):
        folder_path = BASE_DIR / mod["folder"]

        # Module title
        mod_title = doc.add_paragraph()
        run = mod_title.add_run(f"  {idx}.  {mod['display'].upper()}  ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        mod_title.paragraph_format.space_before = Pt(6)

        # Title background via shading on paragraph — use table trick
        title_tbl = doc.add_table(rows=1, cols=1)
        title_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tc = title_tbl.rows[0].cells[0]
        set_cell_bg(tc, "4338CA")
        p = tc.paragraphs[0]
        p.clear()
        r = p.add_run(f"  {idx}.  {mod['display'].upper()}  ")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Description
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.left_indent = Cm(0.5)
        desc_run = desc_p.add_run(f"Description: ")
        desc_run.bold = True
        desc_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        desc_p.add_run(mod["description"]).font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

        # Files table
        file_tbl = doc.add_table(rows=1, cols=3)
        file_tbl.style = "Table Grid"
        fh = file_tbl.rows[0].cells
        for cell, txt in zip(fh, ["#", "Test File", "Functions / Test Methods"]):
            cell.text = txt
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(cell, "6366F1")

        # Set column widths
        for row in file_tbl.rows:
            row.cells[0].width = Cm(1.2)
            row.cells[1].width = Cm(5.5)
            row.cells[2].width = Cm(11.0)

        file_idx = 1
        for fname in mod["files"]:
            fpath = folder_path / fname
            if not fpath.exists():
                continue
            funcs = _extract_functions_from_file(str(fpath))
            public_funcs = [(n, d) for n, d in funcs if not n.startswith("_")]

            func_lines = []
            for fname_f, fdoc in public_funcs:
                if fdoc:
                    func_lines.append(f"• {fname_f}() — {fdoc[:120]}")
                else:
                    func_lines.append(f"• {fname_f}()")

            row_cells = file_tbl.add_row().cells
            row_cells[0].text = str(file_idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = fname
            row_cells[1].paragraphs[0].runs[0].bold = True
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(9)
            row_cells[2].text = "\n".join(func_lines) if func_lines else "—"
            row_cells[2].paragraphs[0].runs[0].font.size = Pt(8.5)
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)

            bg = "F8FAFF" if file_idx % 2 == 0 else "FFFFFF"
            for c in row_cells:
                set_cell_bg(c, bg)
                set_cell_border(c)
            file_idx += 1

        doc.add_paragraph()

    # ── FOOTER NOTE ──────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "This document is auto-generated by the Sparqla QA Automation Framework. "
        "All modules listed above are covered by Selenium-based end-to-end automation tests."
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    fr.italic = True

    doc.save(str(OUTPUT_PATH))
    print(f"✅ Word document generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_word_doc()
